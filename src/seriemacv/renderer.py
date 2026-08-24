"""Deterministic resume projections from canonical career data."""

from __future__ import annotations

import os
import tempfile
from dataclasses import dataclass
from html import escape
from io import BytesIO
from pathlib import Path
from typing import Any, Literal, Protocol

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor

from seriemacv.career import CareerDocument, Education, Experience, Skill
from seriemacv.i18n import translate
from seriemacv.styles import (
    ResumeStyleId,
    StyleManifest,
    StylePackage,
    load_style,
    resolve_resume_color,
)

ResumeLocale = str
ResumeFormat = Literal["markdown", "html", "pdf", "docx"]
_FILENAMES = {
    "markdown": "md",
    "html": "html",
    "pdf": "pdf",
    "docx": "docx",
}


class ResumeRenderError(ValueError):
    pass


@dataclass(frozen=True)
class ResumePresentation:
    career: CareerDocument
    labels: dict[str, str]
    contacts: tuple[str, ...]
    links: tuple[tuple[str, str], ...]
    experience: list[Experience]
    education: list[Education]


class PdfRenderer(Protocol):
    def render(self, html: str) -> bytes: ...


class PlaywrightPdfRenderer:
    def render(self, html: str) -> bytes:
        try:
            from playwright.sync_api import sync_playwright

            with sync_playwright() as playwright:
                browser = playwright.chromium.launch()
                try:
                    page = browser.new_page()
                    page.set_content(html, wait_until="load")
                    return page.pdf(
                        format="A4", print_background=True, prefer_css_page_size=True
                    )
                finally:
                    browser.close()
        except Exception as error:
            raise ResumeRenderError(
                "PDF rendering requires Chromium; run "
                "'python -m playwright install chromium'"
            ) from error


def render_markdown(
    career: CareerDocument,
    locale: ResumeLocale,
    style_id: ResumeStyleId = "clean",
) -> str:
    presentation = _presentation(career, locale)
    style = load_style(style_id).manifest
    _ensure_supported(style, "markdown")
    variant = style.markdown_variant
    show_dividers = style.tokens.section_divider != "none"
    profile = career.profile
    blocks = _markdown_header(presentation, variant, show_dividers)
    if career.summary:
        blocks.append(
            _md_section(
                presentation.labels["summary"],
                career.summary,
                variant,
                show_dividers,
            )
        )
    if presentation.experience:
        blocks.append(
            _md_records(
                presentation.labels["experience"],
                presentation.experience,
                presentation.labels,
                True,
                variant,
                show_dividers,
            )
        )
    if presentation.education:
        blocks.append(
            _md_records(
                presentation.labels["education"],
                presentation.education,
                presentation.labels,
                False,
                variant,
                show_dividers,
            )
        )
    if career.skills:
        blocks.append(
            _md_section(
                presentation.labels["skills"],
                _markdown_skills(career.skills, presentation.labels, variant),
                variant,
                show_dividers,
            )
        )
    if profile.languages:
        language_text = (
            " | ".join(profile.languages)
            if variant == "compact"
            else "\n".join(f"- {item}" for item in profile.languages)
        )
        blocks.append(
            _md_section(
                presentation.labels["languages"],
                language_text,
                variant,
                show_dividers,
            )
        )
    separator = "\n" if variant == "compact" else "\n\n"
    return separator.join(block for block in blocks if block) + "\n"


def render_html(
    career: CareerDocument,
    locale: ResumeLocale,
    style_id: ResumeStyleId = "clean",
    resume_color: str | None = None,
) -> str:
    presentation = _presentation(career, locale)
    style = resolve_resume_color(load_style(style_id), resume_color)
    _ensure_supported(style.manifest, "html")
    header, main, sidebar = _html_parts(presentation, style.manifest)
    return (
        style.template.replace("{{locale}}", locale)
        .replace("{{title}}", escape(career.profile.name))
        .replace("{{css}}", _style_css(style))
        .replace("{{header}}", header)
        .replace("{{main}}", main)
        .replace("{{sidebar}}", sidebar)
        .replace("{{style_id}}", style.manifest.id)
        .replace("{{ats_safe}}", str(style.manifest.ats_safe).lower())
    )


def render_docx(
    career: CareerDocument,
    locale: ResumeLocale,
    style_id: ResumeStyleId = "clean",
    resume_color: str | None = None,
) -> bytes:
    presentation = _presentation(career, locale)
    base_style = load_style(style_id).manifest
    style = resolve_resume_color(load_style(style_id), resume_color).manifest
    _ensure_supported(style, "docx")
    document = Document()
    body_color = (
        base_style.tokens.primary_color
        if resume_color and style.id.startswith(("modern", "clean-executive", "sidebar"))
        else None
    )
    _configure_docx(document, style, body_color=body_color)
    if style.layout == "timeline":
        _docx_timeline_layout(document, presentation, style)
    elif style.id.startswith("left-rail"):
        _docx_left_rail_layout(document, presentation, style)
    else:
        _docx_header(
            document,
            presentation,
            style,
            include_contacts=style.ats_safe or style.id.startswith("contact-band"),
        )
    if style.layout == "two-column" and not style.id.startswith("left-rail"):
        _docx_sidebar_layout(document, presentation, style)
    elif style.layout == "single-column":
        _docx_main_sections(document, presentation, style, include_secondary=True)
    stream = BytesIO()
    document.save(stream)
    return stream.getvalue()


def write_resume(
    project_path: Path,
    career: CareerDocument,
    locale: ResumeLocale,
    output_format: ResumeFormat,
    pdf_renderer: PdfRenderer | None = None,
    style_id: ResumeStyleId = "clean",
    resume_color: str | None = None,
) -> Path:
    path = project_path / "exports" / f"resume.{locale}.{_FILENAMES[output_format]}"
    path.parent.mkdir(parents=True, exist_ok=True)
    style = load_style(style_id).manifest
    _ensure_supported(style, output_format)
    content: bytes
    if output_format == "markdown":
        content = render_markdown(career, locale, style_id).encode()
    elif output_format == "html":
        content = render_html(career, locale, style_id, resume_color).encode()
    elif output_format == "pdf":
        content = (pdf_renderer or PlaywrightPdfRenderer()).render(
            render_html(career, locale, style_id, resume_color)
        )
    else:
        content = render_docx(career, locale, style_id, resume_color)
    _atomic_write(path, content)
    return path


def write_markdown_resume(
    project_path: Path,
    career: CareerDocument,
    locale: ResumeLocale,
    style_id: ResumeStyleId = "clean",
) -> Path:
    return write_resume(project_path, career, locale, "markdown", style_id=style_id)


def _ensure_supported(style: StyleManifest, output_format: str) -> None:
    if output_format not in style.supported_formats:
        raise ResumeRenderError(
            f"Style '{style.id}' does not support format '{output_format}'"
        )


def _style_css(style: StylePackage) -> str:
    tokens = style.manifest.tokens
    values = {
        "font_family": f'"{tokens.font_family}"',
        "body_size_pt": f"{tokens.body_size_pt:g}",
        "name_size_pt": f"{tokens.name_size_pt:g}",
        "section_size_pt": f"{tokens.section_size_pt:g}",
        "record_size_pt": f"{tokens.record_size_pt:g}",
        "margins_mm": f"{tokens.margins_mm:g}",
        "content_width_mm": f"{210 - (2 * tokens.margins_mm):g}",
        "line_spacing": f"{tokens.line_spacing:g}",
        "section_spacing_pt": f"{tokens.section_spacing_pt:g}",
        "primary_color": f"#{tokens.primary_color.lower()}",
        "accent_color": f"#{tokens.accent_color.lower()}",
        "header_alignment": tokens.header_alignment,
    }
    if tokens.sidebar_width_mm is not None:
        values["sidebar_width_mm"] = f"{tokens.sidebar_width_mm:g}"
        values["sidebar_offset_mm"] = f"{tokens.sidebar_width_mm + 7:g}"
    css = style.css
    for name, value in values.items():
        css = css.replace(f"{{{{{name}}}}}", value)
    if "{{" in css or "}}" in css:
        raise ResumeRenderError(
            f"Style '{style.manifest.id}' contains an unresolved CSS token"
        )
    return css


def _presentation(career: CareerDocument, locale: ResumeLocale) -> ResumePresentation:
    profile = career.profile
    links = {**profile.links}
    if profile.linkedin:
        links.setdefault("LinkedIn", profile.linkedin)
    if profile.portfolio:
        links.setdefault("Portfolio", profile.portfolio)
    labels = dict(career.catalog.labels) if career.catalog else _labels(locale)
    if career.catalog:
        labels["__months"] = "\x1f".join(career.catalog.months)
        labels["__date_format"] = career.catalog.date_format
    return ResumePresentation(
        career=career,
        labels=labels,
        contacts=tuple(
            item for item in (profile.location, profile.email, profile.phone) if item
        ),
        links=tuple(links.items()),
        experience=list(_ordered(career.experience)),
        education=list(_ordered(career.education)),
    )


def _markdown_header(
    presentation: ResumePresentation, variant: str, show_dividers: bool
) -> list[str]:
    profile = presentation.career.profile
    contact = " | ".join(presentation.contacts)
    links = [f"{name}: {url}" for name, url in presentation.links]
    if variant == "classic":
        values = [profile.name, "=" * len(profile.name), f"*{profile.title}*"]
    elif variant == "modern":
        values = [f"# {profile.name}", f"**{profile.title}**"]
        if show_dividers:
            values.append("---")
    elif variant == "compact":
        values = [f"# {profile.name} | {profile.title}"]
    elif variant == "clean-executive":
        values = [f"# {profile.name}", f"### {profile.title}"]
    elif variant == "timeline":
        values = [f"## {profile.name}", f"*{profile.title}*"]
    elif variant == "sidebar":
        values = [f"# {profile.name}", f"*{profile.title}*"]
        if contact:
            values.append(f"> {contact}")
        values.extend(f"> {link}" for link in links)
        return values
    else:
        values = [f"# {profile.name}", profile.title]
    if contact:
        values.append(contact)
    values.extend(links)
    return values


def _md_section(
    title: str, content: str, variant: str, show_dividers: bool
) -> str:
    if variant == "classic":
        heading = f"{title}\n{'-' * len(title)}" if show_dividers else f"## {title}"
        return f"{heading}\n\n{content}"
    divider = "\n---" if show_dividers else ""
    if variant == "compact":
        return f"## {title}{divider}\n{content}"
    return f"## {title}{divider}\n\n{content}"


def _md_records(
    title: str,
    records: list[Experience] | list[Education],
    labels: dict[str, str],
    experience: bool,
    variant: str,
    show_dividers: bool,
) -> str:
    values = []
    for record in records:
        heading = _record_heading(record, experience)
        heading_text = (
            f"**{heading}**  " if variant == "compact" else f"### {heading}"
        )
        values.append(
            "\n".join(
                [
                    heading_text,
                    _details(record, labels, experience),
                    *[f"- {highlight}" for highlight in record.highlights],
                ]
            )
        )
    separator = "\n" if variant == "compact" else "\n\n"
    return _md_section(title, separator.join(values), variant, show_dividers)


def _markdown_skills(
    skills: list[Skill], labels: dict[str, str], variant: str
) -> str:
    separator = "\n" if variant == "compact" else "\n\n"
    return separator.join(
        (f"**{category}:** " if category else "")
        + ", ".join(_markdown_skill(skill, labels) for skill in items)
        for category, items in _skill_groups(skills, labels).items()
    )


def _markdown_skill(skill: Skill, labels: dict[str, str]) -> str:
    name = f"**{skill.name}**" if skill.core else skill.name
    return name + (
        f" ({_level_label(labels, skill.level)})"
        if skill.level
        else ""
    )


def _html_parts(
    presentation: ResumePresentation, style: StyleManifest
) -> tuple[str, str, str]:
    if style.layout == "timeline":
        return _html_timeline_parts(presentation)
    profile = presentation.career.profile
    contact = " | ".join(escape(item) for item in presentation.contacts)
    links = " ".join(
        f'<a href="{escape(url, quote=True)}">{escape(name)}</a>'
        for name, url in presentation.links
    )
    header_parts = [
        f"<h1>{escape(profile.name)}</h1>",
        f'<p class="title">{escape(profile.title)}</p>',
    ]
    if style.layout == "single-column":
        header_parts.extend((f"<p>{contact}</p>", f"<p>{links}</p>"))
    elif style.id.startswith("contact-band"):
        header_parts.append(
            f'<div class="contact-band"><p>{contact}</p><p>{links}</p></div>'
        )

    main_sections = []
    if presentation.career.summary:
        main_sections.append(
            _html_section(
                presentation.labels["summary"],
                f"<p>{escape(presentation.career.summary)}</p>",
            )
        )
    if presentation.experience:
        main_sections.append(
            _html_records(
                presentation.labels["experience"],
                presentation.experience,
                presentation.labels,
                True,
            )
        )
    if presentation.education:
        main_sections.append(
            _html_records(
                presentation.labels["education"],
                presentation.education,
                presentation.labels,
                False,
            )
        )

    secondary = []
    if presentation.career.skills:
        secondary.append(
            _html_section(
                presentation.labels["skills"],
                _html_skills(presentation.career.skills, presentation.labels),
            )
        )
    if profile.languages:
        secondary.append(
            _html_section(
                presentation.labels["languages"], _html_list(profile.languages)
            )
        )
    if style.layout == "single-column":
        main_sections.extend(secondary)
        sidebar = ""
    else:
        compact_link_labels = _uses_compact_link_labels(style)
        sidebar_contact = "".join(
            f"<p>{escape(item)}</p>" for item in presentation.contacts
        ) + "".join(
            f'<p><a href="{escape(url, quote=True)}">'
            f"{escape(name) if compact_link_labels else f'{escape(name)}: {escape(url)}'}</a></p>"
            for name, url in presentation.links
        )
        if style.id.startswith("left-rail"):
            rail_profile = (
                f'<div class="rail-profile"><p class="rail-name">{escape(profile.name)}</p>'
                f'<p class="title">{escape(profile.title)}</p></div>'
            )
            header_parts = []
            sidebar = (
                f'<aside class="sidebar">{rail_profile}'
                f'<div class="contact">{sidebar_contact}</div>{"".join(secondary)}</aside>'
            )
        else:
            sidebar = (
                f'<aside class="sidebar"><div class="contact">{sidebar_contact}</div>'
                f'{"".join(secondary)}</aside>'
            )
    return "".join(header_parts), "".join(main_sections), sidebar


def _uses_compact_link_labels(style: StyleManifest) -> bool:
    return style.id.startswith(
        ("left-rail", "detail-sidebar", "sidebar", "split-header")
    )


def _html_timeline_parts(
    presentation: ResumePresentation,
) -> tuple[str, str, str]:
    profile = presentation.career.profile
    contact = " | ".join(escape(item) for item in presentation.contacts)
    links = " ".join(
        f'<a href="{escape(url, quote=True)}">{escape(name)}</a>'
        for name, url in presentation.links
    )
    header = (
        f"<h1>{escape(profile.name)}</h1>"
        f'<p class="title">{escape(profile.title)}</p>'
        f"<p>{contact}</p><p>{links}</p>"
    )
    sections = []
    if presentation.career.summary:
        sections.append(
            _html_section(
                presentation.labels["summary"],
                f"<p>{escape(presentation.career.summary)}</p>",
            )
        )
    if presentation.experience:
        sections.append(
            _html_timeline_records(
                presentation.labels["experience"],
                presentation.experience,
                presentation.labels,
                True,
            )
        )
    if presentation.education:
        sections.append(
            _html_timeline_records(
                presentation.labels["education"],
                presentation.education,
                presentation.labels,
                False,
            )
        )
    if presentation.career.skills:
        sections.append(
            _html_section(
                presentation.labels["skills"],
                _html_skills(presentation.career.skills, presentation.labels),
            )
        )
    if profile.languages:
        sections.append(
            _html_section(
                presentation.labels["languages"], _html_list(profile.languages)
            )
        )
    return header, "".join(sections), ""


def _html_timeline_records(
    title: str,
    records: list[Experience] | list[Education],
    labels: dict[str, str],
    experience: bool,
) -> str:
    rows = []
    for record in records:
        rows.append(
            '<div class="timeline-record">'
            f'<div class="timeline-date">{escape(_date_range(record, labels))}</div>'
            f"<article><h3>{escape(_record_heading(record, experience))}</h3>"
            f"<p>{escape(_details_without_dates(record, experience))}</p>"
            f"{_html_list(record.highlights) if record.highlights else ''}</article>"
            "</div>"
        )
    return _html_section(title, "".join(rows))


def _html_list(values: list[str]) -> str:
    return "<ul>" + "".join(f"<li>{escape(item)}</li>" for item in values) + "</ul>"


def _html_section(title: str, content: str) -> str:
    return f"<section><h2>{escape(title)}</h2>{content}</section>"


def _html_records(
    title: str,
    records: list[Experience] | list[Education],
    labels: dict[str, str],
    experience: bool,
) -> str:
    articles = []
    for record in records:
        articles.append(
            f"<article><h3>{escape(_record_heading(record, experience))}</h3>"
            f"<p>{escape(_details(record, labels, experience))}</p>"
            f"{_html_list(record.highlights) if record.highlights else ''}</article>"
        )
    return _html_section(title, "".join(articles))


def _html_skills(skills: list[Skill], labels: dict[str, str]) -> str:
    groups = []
    for category, items in _skill_groups(skills, labels).items():
        values = []
        for skill in items:
            name = (
                f"<strong>{escape(skill.name)}</strong>"
                if skill.core
                else escape(skill.name)
            )
            level = (
                _level_label(labels, skill.level)
                if skill.level
                else ""
            )
            values.append(name + (f" ({escape(level)})" if level else ""))
        category_text = f"<strong>{escape(category)}:</strong> " if category else ""
        groups.append(f"<p>{category_text}{', '.join(values)}</p>")
    return "".join(groups)


def _configure_docx(
    document: Document, style: StyleManifest, *, body_color: str | None = None
) -> None:
    tokens = style.tokens
    section = document.sections[0]
    section.page_width = Mm(210)
    section.page_height = Mm(297)
    section.top_margin = Mm(tokens.margins_mm)
    section.bottom_margin = Mm(tokens.margins_mm)
    section.left_margin = Mm(tokens.margins_mm)
    section.right_margin = Mm(tokens.margins_mm)
    normal = document.styles["Normal"]
    normal.font.name = tokens.font_family
    normal.font.size = Pt(tokens.body_size_pt)
    normal.font.color.rgb = RGBColor.from_string(body_color or tokens.primary_color)
    normal.paragraph_format.line_spacing = tokens.line_spacing
    normal.paragraph_format.space_after = Pt(0)
    heading = document.styles.add_style("Resume Heading", WD_STYLE_TYPE.PARAGRAPH)
    heading.font.name = tokens.font_family
    heading.font.size = Pt(tokens.section_size_pt)
    heading.font.bold = True
    heading.font.color.rgb = RGBColor.from_string(tokens.primary_color)
    heading.paragraph_format.space_before = Pt(tokens.section_spacing_pt)
    heading.paragraph_format.space_after = Pt(4)


def _docx_header(
    document: Document,
    presentation: ResumePresentation,
    style: StyleManifest,
    *,
    include_contacts: bool,
    compact_link_labels: bool = False,
) -> None:
    profile = presentation.career.profile
    alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
        if style.tokens.header_alignment == "center"
        else WD_ALIGN_PARAGRAPH.LEFT
    )
    name = document.add_paragraph()
    name.alignment = alignment
    name.paragraph_format.space_after = Pt(2)
    name_run = name.add_run(profile.name)
    name_run.bold = True
    name_run.font.size = Pt(style.tokens.name_size_pt)
    name_run.font.color.rgb = RGBColor.from_string(style.tokens.primary_color)
    title = document.add_paragraph(profile.title)
    title.alignment = alignment
    title.paragraph_format.space_after = Pt(2)
    if include_contacts and presentation.contacts:
        paragraph = document.add_paragraph(" | ".join(presentation.contacts))
        paragraph.alignment = alignment
    if include_contacts:
        for label, url in presentation.links:
            paragraph = document.add_paragraph()
            if compact_link_labels:
                _docx_hyperlink(paragraph, label, url)
            else:
                paragraph.add_run(f"{label}: {url}")
            paragraph.alignment = alignment


def _docx_main_sections(
    container: Any,
    presentation: ResumePresentation,
    style: StyleManifest,
    *,
    include_secondary: bool,
) -> None:
    career = presentation.career
    if career.summary:
        _docx_section(container, presentation.labels["summary"], style)
        container.add_paragraph(career.summary)
    if presentation.experience:
        _docx_records(
            container,
            presentation.labels["experience"],
            presentation.experience,
            presentation.labels,
            True,
            style,
        )
    if presentation.education:
        _docx_records(
            container,
            presentation.labels["education"],
            presentation.education,
            presentation.labels,
            False,
            style,
        )
    if include_secondary:
        _docx_secondary_sections(container, presentation, style)


def _docx_secondary_sections(
    container: Any, presentation: ResumePresentation, style: StyleManifest
) -> None:
    career = presentation.career
    if career.skills:
        _docx_section(container, presentation.labels["skills"], style)
        _docx_skills(container, career.skills, presentation.labels)
    if career.profile.languages:
        _docx_section(container, presentation.labels["languages"], style)
        container.add_paragraph(" | ".join(career.profile.languages))


def _docx_hyperlink(paragraph: Any, label: str, url: str) -> None:
    """Append a visible external hyperlink without exposing its URL as text."""
    relationship_id = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend((color, underline))
    text = OxmlElement("w:t")
    text.text = label
    run.extend((properties, text))
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _docx_sidebar_layout(
    document: Document, presentation: ResumePresentation, style: StyleManifest
) -> None:
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders_none(table)
    main_cell, sidebar_cell = table.rows[0].cells
    available_width = 210 - (2 * style.tokens.margins_mm)
    sidebar_width = style.tokens.sidebar_width_mm or 55
    main_width = available_width - sidebar_width - 7
    table.columns[0].width = Mm(main_width)
    table.columns[1].width = Mm(sidebar_width)
    main_cell.width = Mm(main_width)
    sidebar_cell.width = Mm(sidebar_width)
    main_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    sidebar_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _docx_main_sections(main_cell, presentation, style, include_secondary=False)
    for item in presentation.contacts:
        sidebar_cell.add_paragraph(item)
    for label, url in presentation.links:
        paragraph = sidebar_cell.add_paragraph()
        if _uses_compact_link_labels(style):
            _docx_hyperlink(paragraph, label, url)
        else:
            paragraph.add_run(f"{label}: {url}")
    _docx_secondary_sections(sidebar_cell, presentation, style)
    _remove_empty_leading_paragraph(main_cell)
    _remove_empty_leading_paragraph(sidebar_cell)


def _docx_left_rail_layout(
    document: Document, presentation: ResumePresentation, style: StyleManifest
) -> None:
    """Render the non-ATS left rail with the profile in the colored cell."""
    table = document.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders_none(table)
    rail_cell, main_cell = table.rows[0].cells
    available_width = 210 - (2 * style.tokens.margins_mm)
    rail_width = style.tokens.sidebar_width_mm or 52
    main_width = available_width - rail_width
    table.columns[0].width = Mm(rail_width)
    table.columns[1].width = Mm(main_width)
    rail_cell.width = Mm(rail_width)
    main_cell.width = Mm(main_width)
    rail_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    main_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _set_cell_shading(rail_cell, style.tokens.primary_color)
    _docx_header(
        rail_cell,
        presentation,
        style,
        include_contacts=True,
        compact_link_labels=True,
    )
    _docx_secondary_sections(rail_cell, presentation, style)
    _docx_main_sections(main_cell, presentation, style, include_secondary=False)
    _remove_empty_leading_paragraph(rail_cell)
    _remove_empty_leading_paragraph(main_cell)
    white = RGBColor(255, 255, 255)
    for paragraph in rail_cell.paragraphs:
        for run in paragraph.runs:
            run.font.color.rgb = white


def _docx_timeline_layout(
    document: Document, presentation: ResumePresentation, style: StyleManifest
) -> None:
    table = document.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    _set_table_borders_none(table)

    _, header_cell = _docx_timeline_row(table, style)
    _docx_header(header_cell, presentation, style, include_contacts=True)
    _remove_empty_leading_paragraph(header_cell)

    career = presentation.career
    if career.summary:
        _, content = _docx_timeline_row(table, style)
        _docx_section(content, presentation.labels["summary"], style)
        content.add_paragraph(career.summary)
        _remove_empty_leading_paragraph(content)
    if presentation.experience:
        _docx_timeline_heading_row(
            table, presentation.labels["experience"], style
        )
        _docx_timeline_records(
            table,
            presentation.experience,
            presentation.labels,
            True,
            style,
        )
    if presentation.education:
        _docx_timeline_heading_row(table, presentation.labels["education"], style)
        _docx_timeline_records(
            table,
            presentation.education,
            presentation.labels,
            False,
            style,
        )
    if career.skills:
        _, content = _docx_timeline_row(table, style)
        _docx_section(content, presentation.labels["skills"], style)
        _docx_skills(content, career.skills, presentation.labels)
        _remove_empty_leading_paragraph(content)
    if career.profile.languages:
        _, content = _docx_timeline_row(table, style)
        _docx_section(content, presentation.labels["languages"], style)
        content.add_paragraph(" | ".join(career.profile.languages))
        _remove_empty_leading_paragraph(content)


def _docx_timeline_heading_row(table: Any, title: str, style: StyleManifest) -> None:
    _, content = _docx_timeline_row(table, style)
    _docx_section(content, title, style)
    _remove_empty_leading_paragraph(content)


def _docx_skills(
    container: Any, skills: list[Skill], labels: dict[str, str]
) -> None:
    for category, grouped_skills in _skill_groups(skills, labels).items():
        paragraph = container.add_paragraph()
        if category:
            category_run = paragraph.add_run(f"{category}: ")
            category_run.bold = True
        for index, skill in enumerate(grouped_skills):
            if index:
                paragraph.add_run(", ")
            skill_run = paragraph.add_run(skill.name)
            skill_run.bold = skill.core
            if skill.level:
                paragraph.add_run(
                    f" ({_level_label(labels, skill.level)})"
                )


def _docx_timeline_records(
    table: Any,
    records: list[Experience] | list[Education],
    labels: dict[str, str],
    experience: bool,
    style: StyleManifest,
) -> None:
    white = RGBColor(255, 255, 255)
    for record in records:
        date_cell, content = _docx_timeline_row(table, style)
        date = date_cell.paragraphs[0]
        date.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date.add_run(_date_range(record, labels))
        date_run.font.color.rgb = white
        name = content.add_paragraph()
        name_run = name.add_run(_record_heading(record, experience))
        name_run.bold = True
        name_run.font.size = Pt(style.tokens.record_size_pt)
        details = _details_without_dates(record, experience)
        if details:
            content.add_paragraph(details)
        for highlight in record.highlights:
            _docx_plain_bullet(content, highlight)
        _remove_empty_leading_paragraph(content)


def _docx_timeline_row(table: Any, style: StyleManifest) -> tuple[Any, Any]:
    date_cell, content_cell = table.add_row().cells
    available_width = 210 - (2 * style.tokens.margins_mm)
    rail_width = style.tokens.sidebar_width_mm or 36
    content_width = available_width - rail_width
    table.columns[0].width = Mm(rail_width)
    table.columns[1].width = Mm(content_width)
    date_cell.width = Mm(rail_width)
    content_cell.width = Mm(content_width)
    date_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    content_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    _set_cell_shading(date_cell, style.tokens.primary_color)
    return date_cell, content_cell


def _docx_section(container: Any, title: str, style: StyleManifest) -> None:
    paragraph = container.add_paragraph(title, style="Resume Heading")
    if style.tokens.section_divider == "none":
        return
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(
        qn("w:sz"), "8" if style.tokens.section_divider == "accent" else "4"
    )
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), style.tokens.accent_color)
    borders.append(bottom)
    paragraph._p.get_or_add_pPr().append(borders)


def _docx_records(
    container: Any,
    title: str,
    records: list[Experience] | list[Education],
    labels: dict[str, str],
    experience: bool,
    style: StyleManifest,
) -> None:
    _docx_section(container, title, style)
    for record in records:
        name = container.add_paragraph()
        name_run = name.add_run(_record_heading(record, experience))
        name_run.bold = True
        name_run.font.size = Pt(style.tokens.record_size_pt)
        container.add_paragraph(_details(record, labels, experience))
        for highlight in record.highlights:
            _docx_plain_bullet(container, highlight)


def _docx_plain_bullet(container: Any, value: str) -> None:
    """Use visible text rather than Word list XML for straightforward ATS parsing."""
    paragraph = container.add_paragraph(f"- {value}")
    paragraph.paragraph_format.left_indent = Mm(4.5)
    paragraph.paragraph_format.first_line_indent = Mm(-3)


def _set_table_borders_none(table: Any) -> None:
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "nil")
        borders.append(element)
    table._tbl.tblPr.append(borders)


def _set_cell_shading(cell: Any, color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    cell._tc.get_or_add_tcPr().append(shading)


def _remove_empty_leading_paragraph(cell: Any) -> None:
    paragraph = cell.paragraphs[0]
    if not paragraph.text and len(cell.paragraphs) > 1:
        paragraph._element.getparent().remove(paragraph._element)


def _record_heading(record: Experience | Education, experience: bool) -> str:
    values = (
        (record.title, record.company)
        if experience
        else (record.degree, record.institution)
    )
    return " - ".join(values)


def _labels(locale: str) -> dict[str, str]:
    if locale not in {"pt-BR", "en"}:
        raise ValueError(f"Unsupported resume locale: {locale}")
    return {
        key: translate(locale, key)
        for key in (
            "summary",
            "experience",
            "education",
            "skills",
            "languages",
            "current",
            "other",
            "level.beginner",
            "level.intermediate",
            "level.advanced",
            "level.expert",
        )
    }


def _ordered(
    records: list[Experience] | list[Education],
) -> list[Experience] | list[Education]:
    return sorted(
        records,
        key=lambda record: (
            record.end_date is None,
            record.end_date or "",
            record.start_date,
        ),
        reverse=True,
    )


def _details(
    record: Experience | Education, labels: dict[str, str], experience: bool
) -> str:
    values = [_details_without_dates(record, experience)]
    values.append(_date_range(record, labels))
    return " | ".join(value for value in values if value)


def _details_without_dates(
    record: Experience | Education, experience: bool
) -> str:
    values = (
        [record.location] if experience else [record.field_of_study, record.location]
    )
    if experience and record.employment_type:
        values.append(record.employment_type)
    return " | ".join(value for value in values if value)


def _date_range(
    record: Experience | Education, labels: dict[str, str]
) -> str:
    end = (
        _format_date(record.end_date, labels) if record.end_date else labels["current"]
    )
    return " - ".join((_format_date(record.start_date, labels), end))


def _skill_groups(
    skills: list[Skill], labels: dict[str, str]
) -> dict[str, list[Skill]]:
    groups: dict[str, list[Skill]] = {}
    has_category = any(skill.category for skill in skills)
    for skill in skills:
        category = skill.category or (labels["other"] if has_category else "")
        groups.setdefault(category, []).append(skill)
    return groups


def _level_label(labels: dict[str, str], level: str) -> str:
    return labels[f"level.{level}"]


def _format_date(value: str, labels: dict[str, str]) -> str:
    year, month = value.split("-")
    if "__months" in labels:
        name = labels["__months"].split("\x1f")[int(month) - 1]
        return labels["__date_format"].replace("{month}", name).replace("{year}", year)
    # Localized documents carry their catalog; built-in locales retain the
    # previous deterministic formatting path for renderer-only callers.
    names = None
    if labels["current"] == "Atual":
        names = ("Jan.", "Fev.", "Mar.", "Abr.", "Mai.", "Jun.", "Jul.", "Ago.", "Set.", "Out.", "Nov.", "Dez.")
        return f"{names[int(month) - 1]} de {year}"
    return f"{('Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun', 'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec')[int(month) - 1]} {year}"


def _atomic_write(path: Path, content: bytes) -> None:
    descriptor, temporary = tempfile.mkstemp(
        dir=path.parent, prefix=f".{path.name}.", suffix=".tmp"
    )
    try:
        with os.fdopen(descriptor, "wb") as file:
            file.write(content)
        os.replace(temporary, path)
    except BaseException:
        Path(temporary).unlink(missing_ok=True)
        raise
