"""Optional, local NuExtract resume-import proposals."""
from __future__ import annotations

import base64
import hashlib
import json
import urllib.error
import urllib.request
from datetime import UTC, datetime
from pathlib import Path
from typing import Any, Literal
from urllib.parse import urlparse

from docx import Document
from pydantic import BaseModel, ConfigDict, Field, ValidationError, field_validator
from ruamel.yaml import YAML

from seriemacv.career import CAREER_FILE, CareerFactsDocument, CareerLocaleDocument, _atomic_write, locale_path

PROPOSALS_DIRECTORY = "proposals"
MAX_SOURCE_BYTES = 10 * 1024 * 1024


class ImportError(ValueError):
    pass


class StrictModel(BaseModel):
    model_config = ConfigDict(extra="forbid", strict=True)


class NuExtractSettings(StrictModel):
    endpoint: str
    model: str = Field(min_length=1)
    timeout_seconds: int = Field(default=60, ge=1, le=600)
    max_output_tokens: int = Field(default=8192, ge=256, le=32768)
    multimodal: bool = False

    @field_validator("endpoint")
    @classmethod
    def local_endpoint(cls, value: str) -> str:
        parsed = urlparse(value)
        if parsed.scheme not in {"http", "https"} or parsed.hostname not in {"127.0.0.1", "localhost", "::1"}:
            raise ValueError("must be an HTTP(S) loopback endpoint")
        return value.rstrip("/")


class SourceExcerpt(StrictModel):
    text: str = Field(min_length=1, max_length=4000)
    page: int | None = Field(default=None, ge=1)
    section: str = ""


class ImportProposal(StrictModel):
    schema_version: Literal[1]
    id: str = Field(pattern=r"^import-[a-z0-9-]+$")
    created_at: str
    source_path: str
    source_sha256: str
    language: str
    runtime: NuExtractSettings
    confidence: float = Field(ge=0, le=1)
    pending: list[str] = Field(default_factory=list)
    excerpts: list[SourceExcerpt] = Field(min_length=1)
    career: CareerFactsDocument
    locale: CareerLocaleDocument


def prepare_document(path: Path, *, multimodal: bool) -> tuple[str, list[dict[str, str]], list[SourceExcerpt]]:
    source = path.expanduser().resolve()
    if not source.is_file() or source.stat().st_size > MAX_SOURCE_BYTES:
        raise ImportError("source must be an existing file no larger than 10 MiB")
    suffix = source.suffix.lower()
    if suffix in {".md", ".markdown"}:
        text = source.read_text(encoding="utf-8")
        return text, [], [SourceExcerpt(text=text[:4000], section="markdown")]
    if suffix == ".docx":
        document = Document(source)
        text = "\n".join(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
        for table in document.tables:
            text += "\n" + "\n".join(" | ".join(cell.text for cell in row.cells) for row in table.rows)
        return text, [], [SourceExcerpt(text=text[:4000], section="docx")]
    if suffix != ".pdf":
        raise ImportError("supported source formats are Markdown, DOCX and PDF")
    try:
        import fitz  # type: ignore[import-not-found]
    except ImportError as error:
        raise ImportError("PDF import requires 'pip install .[import]'") from error
    document = fitz.open(source)
    pages = [(index + 1, page.get_text().strip()) for index, page in enumerate(document)]
    text = "\n".join(value for _, value in pages)
    excerpts = [SourceExcerpt(text=value[:4000], page=number) for number, value in pages if value]
    images: list[dict[str, str]] = []
    if any(not value for _, value in pages):
        if not multimodal:
            raise ImportError("PDF contains visual pages; configure a multimodal NuExtract endpoint")
        for number, value in pages:
            if not value:
                pixmap = document[number - 1].get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
                images.append({"page": str(number), "data": base64.b64encode(pixmap.tobytes("png")).decode()})
    if not text and not images:
        raise ImportError("PDF contains no extractable text or images")
    return text, images, excerpts or [SourceExcerpt(text="visual PDF pages", section="pdf")]


def request_proposal(settings: NuExtractSettings, *, text: str, images: list[dict[str, str]], language: str) -> dict[str, Any]:
    template = ImportProposal.model_json_schema()
    content: list[dict[str, Any]] = [{"type": "text", "text": f"Extract only verified resume facts in {language}. Return JSON matching this schema: {json.dumps(template)}\nDocument:\n{text}"}]
    content.extend({"type": "image_url", "image_url": {"url": f"data:image/png;base64,{image['data']}"}} for image in images)
    body = json.dumps({"model": settings.model, "temperature": 0, "max_tokens": settings.max_output_tokens, "response_format": {"type": "json_object"}, "messages": [{"role": "user", "content": content}]}).encode()
    request = urllib.request.Request(f"{settings.endpoint}/v1/chat/completions", data=body, headers={"Content-Type": "application/json"})
    try:
        with urllib.request.urlopen(request, timeout=settings.timeout_seconds) as response:
            payload = json.load(response)
    except (urllib.error.URLError, TimeoutError, OSError) as error:
        raise ImportError(f"NuExtract endpoint failed: {error}") from error
    try:
        return json.loads(payload["choices"][0]["message"]["content"])
    except (KeyError, IndexError, TypeError, json.JSONDecodeError) as error:
        raise ImportError("NuExtract returned invalid JSON") from error


def save_proposal(project_path: Path, proposal: ImportProposal) -> Path:
    directory = project_path / PROPOSALS_DIRECTORY
    directory.mkdir(exist_ok=True)
    path = directory / f"{proposal.id}.yml"
    if path.exists():
        raise ImportError(f"proposal already exists: {proposal.id}")
    yaml = YAML()
    with path.open("w", encoding="utf-8") as stream:
        yaml.dump(proposal.model_dump(mode="json"), stream)
    return path


def load_proposal(project_path: Path, proposal_id: str) -> ImportProposal:
    path = project_path / PROPOSALS_DIRECTORY / f"{proposal_id}.yml"
    if not path.is_file():
        raise ImportError(f"proposal is missing: {proposal_id}")
    return ImportProposal.model_validate(YAML(typ="safe").load(path.read_text(encoding="utf-8")))


def new_proposal(project_path: Path, source: Path, language: str, settings: NuExtractSettings) -> ImportProposal:
    text, images, excerpts = prepare_document(source, multimodal=settings.multimodal)
    payload = request_proposal(settings, text=text, images=images, language=language)
    now = datetime.now(UTC).strftime("%Y%m%d%H%M%S")
    payload.update({"schema_version": 1, "id": f"import-{now}", "created_at": datetime.now(UTC).isoformat(), "source_path": str(source.expanduser().resolve()), "source_sha256": hashlib.sha256(source.read_bytes()).hexdigest(), "language": language, "runtime": settings.model_dump(), "excerpts": [item.model_dump() for item in excerpts]})
    try:
        proposal = ImportProposal.model_validate(payload)
    except ValidationError as error:
        raise ImportError(f"NuExtract proposal is invalid: {error}") from error
    if proposal.locale.locale != language:
        raise ImportError("proposal locale does not match --language")
    return proposal


def apply_proposal(project_path: Path, proposal_id: str) -> tuple[Path, Path]:
    """Apply only to an empty project, preserving existing career facts."""
    proposal = load_proposal(project_path, proposal_id)
    career_path = project_path / CAREER_FILE
    facts = CareerFactsDocument.model_validate(YAML(typ="safe").load(career_path.read_text(encoding="utf-8")))
    empty = CareerFactsDocument(schema_version=2)
    if facts != empty:
        raise ImportError("career.yml already contains data; resolve the proposal manually")
    translated_path = locale_path(project_path, proposal.language)
    if translated_path.exists():
        existing = CareerLocaleDocument.model_validate(YAML(typ="safe").load(translated_path.read_text(encoding="utf-8")))
        if existing.profile.title or existing.summary or existing.experience or existing.education or existing.skills:
            raise ImportError("target locale already contains data; resolve the proposal manually")
    yaml = YAML()
    from io import StringIO
    facts_stream, locale_stream = StringIO(), StringIO()
    yaml.dump(proposal.career.model_dump(mode="json"), facts_stream)
    yaml.dump(proposal.locale.model_dump(mode="json"), locale_stream)
    _atomic_write(career_path, facts_stream.getvalue())
    _atomic_write(translated_path, locale_stream.getvalue())
    return career_path, translated_path
