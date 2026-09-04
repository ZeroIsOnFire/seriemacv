from __future__ import annotations

import tempfile
import unittest
from contextlib import redirect_stderr, redirect_stdout
from io import BytesIO, StringIO
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx.shared import Pt, RGBColor

from seriemacv.career import CareerDocument
from seriemacv.cli import main
from seriemacv.project import create_project
from seriemacv.renderer import (
    render_docx,
    render_html,
    render_markdown,
    write_markdown_resume,
    write_resume,
)
from seriemacv.styles import STYLE_FAMILIES, STYLE_IDS, load_style


def _docx_all_paragraphs(document: Document) -> list[object]:
    paragraphs = list(document.paragraphs)
    for table in document.tables:
        paragraphs.extend(
            paragraph
            for row in table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
        )
    return paragraphs


def _docx_resume_headings(document: Document) -> list[object]:
    return [
        paragraph
        for paragraph in _docx_all_paragraphs(document)
        if paragraph.style.name == "Resume Heading"
    ]


class MarkdownRendererTests(unittest.TestCase):
    def test_all_styles_render_markdown_html_and_docx_in_both_locales(self) -> None:
        for style_id in STYLE_IDS:
            for locale in ("pt-BR", "en"):
                with self.subTest(style=style_id, locale=locale):
                    markdown = render_markdown(self._career(), locale, style_id)
                    html = render_html(self._career(), locale, style_id)
                    document = Document(BytesIO(render_docx(self._career(), locale, style_id)))
                    highlight_label = (
                        "Destaque"
                        if locale == "pt-BR"
                        else "Highlight"
                    )
                    highlight = f"{highlight_label}: Built reliable tools."

                    self.assertIn("Seriema Example", markdown)
                    self.assertIn(f"**{highlight_label}:** Built reliable tools.", markdown)
                    self.assertIn(f'data-style="{style_id}"', html)
                    self.assertIn(
                        f"<strong>{highlight_label}:</strong> Built reliable tools.", html
                    )
                    self.assertNotIn("{{", html)
                    self.assertIn(
                        "Seriema Example",
                        "\n".join(p.text for p in _docx_all_paragraphs(document)),
                    )
                    self.assertIn(
                        highlight,
                        "\n".join(p.text for p in _docx_all_paragraphs(document)),
                    )
                    self.assertEqual(len(document.inline_shapes), 0)
                    manifest = load_style(style_id).manifest
                    if manifest.layout == "two-column":
                        self.assertIn('<aside class="sidebar">', html)
                        self.assertEqual(len(document.tables), 1)
                        table_text = "\n".join(
                            paragraph.text
                            for row in document.tables[0].rows
                            for cell in row.cells
                            for paragraph in cell.paragraphs
                        )
                        expected_heading = (
                            "Experiência profissional"
                            if locale == "pt-BR"
                            else "Professional Experience"
                        )
                        self.assertIn(expected_heading, table_text)
                        if style_id.startswith(
                            ("left-rail", "detail-sidebar", "sidebar", "split-header")
                        ):
                            self.assertIn("Portfolio", table_text)
                            self.assertNotIn("https://example.invalid/seriema", table_text)
                        else:
                            self.assertIn(
                                "Portfolio: https://example.invalid/seriema", table_text
                            )
                        self.assertIn('w:val="nil"', document.tables[0]._tbl.xml)
                    elif manifest.layout == "timeline":
                        self.assertIn('class="timeline-record"', html)
                        self.assertIn('class="timeline-date"', html)
                        self.assertNotIn("<img", html)
                        self.assertNotIn("<svg", html)
                        self.assertGreater(len(document.tables), 0)
                        self.assertIn('w:val="nil"', document.tables[0]._tbl.xml)
                    else:
                        self.assertNotIn('<aside class="sidebar">', html)
                        self.assertEqual(document.tables, [])
                        self.assertNotIn(
                            "List Bullet",
                            [paragraph.style.name for paragraph in document.paragraphs],
                        )

    def test_markdown_styles_have_distinct_structures(self) -> None:
        rendered = {
            style_id: render_markdown(self._career(), "en", style_id)
            for style_id in STYLE_IDS
        }

        self.assertIn("=============", rendered["classic"])
        self.assertIn("# Seriema Example | Software Engineer", rendered["compact"])
        self.assertIn("> Remote | seriema@example.invalid", rendered["sidebar"])
        self.assertIn("### Software Engineer", rendered["clean-executive"])
        self.assertIn("## Seriema Example", rendered["timeline"])

    def test_standard_and_alt_styles_toggle_section_dividers(self) -> None:
        for family in STYLE_FAMILIES:
            with self.subTest(family=family):
                standard_html = render_html(self._career(), "en", family)
                alt_html = render_html(self._career(), "en", f"{family}-alt")
                standard_docx = Document(
                    BytesIO(render_docx(self._career(), "en", family))
                )
                alt_docx = Document(
                    BytesIO(render_docx(self._career(), "en", f"{family}-alt"))
                )
                standard_headings = _docx_resume_headings(standard_docx)
                alt_headings = _docx_resume_headings(alt_docx)

                self.assertTrue(standard_headings)
                self.assertTrue(alt_headings)
                self.assertIn("border-bottom", standard_html)
                self.assertNotIn("border-bottom", alt_html)
                self.assertTrue(
                    all("w:pBdr" in paragraph._p.xml for paragraph in standard_headings)
                )
                self.assertTrue(
                    all("w:pBdr" not in paragraph._p.xml for paragraph in alt_headings)
                )
                self.assertIn("\n---", render_markdown(self._career(), "en", family))
                self.assertNotIn(
                    "\n---", render_markdown(self._career(), "en", f"{family}-alt")
                )

    def test_configurable_styles_apply_one_color_to_html_and_docx(self) -> None:
        for style_id in STYLE_IDS:
            if not load_style(style_id).manifest.color_customizable:
                continue
            with self.subTest(style=style_id):
                html = render_html(self._career(), "en", style_id, "Aa12Bc")
                document = Document(BytesIO(render_docx(self._career(), "en", style_id, "AA12BC")))

                self.assertIn("#aa12bc", html)
                self.assertIn('AA12BC', document.element.xml)

    def test_modern_and_clean_executive_keep_body_text_neutral(self) -> None:
        for style_id, body_color in (
            ("modern", "17324D"),
            ("modern-alt", "17324D"),
            ("clean-executive", "243238"),
            ("clean-executive-alt", "243238"),
            ("sidebar", "17324D"),
            ("sidebar-alt", "17324D"),
        ):
            with self.subTest(style=style_id):
                html = render_html(self._career(), "en", style_id, "AA12BC")
                document = Document(
                    BytesIO(render_docx(self._career(), "en", style_id, "AA12BC"))
                )

                self.assertIn(f"body {{ margin: 0; color: #{body_color.lower()}", html)
                self.assertEqual(
                    document.styles["Normal"].font.color.rgb,
                    RGBColor.from_string(body_color),
                )
                self.assertIn("AA12BC", document.element.xml)

    def test_sidebar_uses_project_color_for_headings_not_body_text(self) -> None:
        html = render_html(self._career(), "en", "sidebar", "AA12BC")

        self.assertIn("body { margin: 0; color: #17324d", html)
        self.assertIn("h1 { margin: 0 0 2pt; color: #aa12bc", html)
        self.assertIn("h2 { margin:", html)
        self.assertIn("color: #aa12bc", html)
        self.assertIn(".sidebar h2 { border-color: #fff; color: #fff; }", html)

    def test_visual_detail_styles_render_contact_values_without_titles(self) -> None:
        for style_id in ("left-rail", "detail-sidebar", "sidebar", "split-header"):
            with self.subTest(style=style_id):
                html = render_html(self._career(), "en", style_id)
                document = Document(
                    BytesIO(render_docx(self._career(), "en", style_id))
                )

                self.assertNotIn("Location: ", html)
                self.assertNotIn("Email: ", html)
                self.assertNotIn("Phone: ", html)
                self.assertIn("Remote", html)
                self.assertIn("seriema@example.invalid", html)
                self.assertIn("+1 555 0100", html)
                document_text = "\n".join(
                    paragraph.text for paragraph in _docx_all_paragraphs(document)
                )
                self.assertNotIn("Location: Remote", document_text)
                self.assertNotIn("Email: seriema@example.invalid", document_text)
                self.assertNotIn("Phone: +1 555 0100", document_text)
                self.assertIn("Remote", document_text)
                self.assertIn("seriema@example.invalid", document_text)
                self.assertIn("+1 555 0100", document_text)

    def test_visual_detail_styles_show_labeled_clickable_links(self) -> None:
        for style_id in ("left-rail", "detail-sidebar", "sidebar", "split-header"):
            with self.subTest(style=style_id):
                html = render_html(self._career(), "en", style_id)
                document = Document(
                    BytesIO(render_docx(self._career(), "en", style_id))
                )

                self.assertIn(
                    '<a href="https://example.invalid/seriema">Portfolio</a>', html
                )
                self.assertNotIn("Portfolio: https://example.invalid/seriema", html)
                self.assertIn("Portfolio", document.element.xml)
                self.assertNotIn(
                    "Portfolio: https://example.invalid/seriema", document.element.xml
                )
                self.assertIn(
                    "https://example.invalid/seriema",
                    "\n".join(
                        relationship.target_ref
                        for relationship in document.part.rels.values()
                        if relationship.is_external
                    ),
                )

    def test_full_bleed_styles_keep_vertical_margins_on_every_page(self) -> None:
        for style_id in (
            "contact-band",
            "contact-band-alt",
            "detail-sidebar",
            "detail-sidebar-alt",
        ):
            with self.subTest(style=style_id):
                html = render_html(self._career(), "en", style_id)

                self.assertIn("@page:first { margin-top: 0; }", html)
                self.assertNotIn("@page { size: A4; margin: 0; }", html)
                self.assertNotIn("margin: -", html)
        for style_id in ("left-rail", "left-rail-alt", "timeline", "timeline-alt"):
            with self.subTest(style=style_id):
                html = render_html(self._career(), "en", style_id)

                self.assertIn("@page { size: A4; margin: 0; }", html)
                self.assertIn('class="page-frame" role="presentation"', html)
                self.assertIn(".page-frame > thead, .page-frame > tfoot { height: 14mm; }", html)
                self.assertNotIn("@page:first", html)
                self.assertNotIn("margin: -", html)
        contact_band = render_html(self._career(), "en", "contact-band")
        self.assertIn("@page { size: A4; margin: 12mm 0; }", contact_band)
        self.assertIn("header { background:", contact_band)
        timeline = render_html(self._career(), "en", "timeline")
        self.assertIn("@page { size: A4; margin: 0; }", timeline)
        self.assertNotIn("section, article { break-inside: avoid; }", timeline)
        self.assertIn("margin: 0 0 7pt -44mm; break-inside: avoid;", timeline)

    def test_colored_side_rails_bleed_through_vertical_page_margins(self) -> None:
        for style_id in ("sidebar", "sidebar-alt"):
            with self.subTest(style=style_id):
                html = render_html(self._career(), "en", style_id)
                self.assertIn("body::before { position: fixed;", html)
                self.assertIn(
                    ".page-frame > thead > tr > td, .page-frame > tfoot > tr > td { height: 12mm; }",
                    html,
                )
                self.assertNotIn("@top-right", html)
                self.assertNotIn("@bottom-right", html)

        for style_id in ("left-rail", "left-rail-alt"):
            with self.subTest(style=style_id):
                html = render_html(self._career(), "en", style_id)
                self.assertIn("body::before { position: fixed;", html)
                self.assertNotIn("@top-left", html)
                self.assertNotIn("@bottom-left", html)

        for style_id in ("timeline", "timeline-alt"):
            with self.subTest(style=style_id):
                html = render_html(self._career(), "en", style_id)
                self.assertIn(
                    ".date-rail { position: fixed; inset: 0 auto 0 0;",
                    html,
                )
                self.assertNotIn("@top-left", html)
                self.assertNotIn("@bottom-left", html)

    def test_non_configurable_styles_and_markdown_ignore_resume_color(self) -> None:
        html = render_html(self._career(), "en", "clean", "AA12BC")

        self.assertNotIn("#aa12bc", html)
        with tempfile.TemporaryDirectory() as temporary_directory:
            output_path = write_resume(
                Path(temporary_directory),
                self._career(),
                "en",
                "markdown",
                style_id="modern",
                resume_color="AA12BC",
            )
            self.assertEqual(
                output_path.read_text(encoding="utf-8"),
                render_markdown(self._career(), "en", "modern"),
            )

    def test_classic_never_draws_a_header_divider(self) -> None:
        for style_id in ("classic", "classic-alt"):
            with self.subTest(style=style_id):
                html = render_html(self._career(), "en", style_id)
                header_rule = html.split("header {", 1)[1].split("}", 1)[0]

                self.assertNotIn("border", header_rule)

    def test_compact_keeps_contact_information_below_the_full_width_name(self) -> None:
        for style_id in ("compact", "compact-alt"):
            with self.subTest(style=style_id):
                html = render_html(self._career(), "en", style_id)
                header_rule = html.split("header {", 1)[1].split("}", 1)[0]

                self.assertIn("display: block", header_rule)
                self.assertNotIn("grid-template-columns", header_rule)

    def test_timeline_is_photo_free_and_places_each_date_in_the_rail(self) -> None:
        html = render_html(self._career(), "en", "timeline")
        document = Document(BytesIO(render_docx(self._career(), "en", "timeline")))
        table_text = "\n".join(
            paragraph.text
            for row in document.tables[0].rows
            for cell in row.cells
            for paragraph in cell.paragraphs
        )

        self.assertIn(
            '<div class="timeline-date">Jan 2024 - Present</div>', html
        )
        self.assertEqual(html.count("Jan 2024 - Present"), 1)
        self.assertNotIn("<img", html)
        self.assertNotIn("<svg", html)
        self.assertNotIn("overflow: hidden", html)
        self.assertIn("Jan 2024 - Present", table_text)
        self.assertIn("Software Engineer - Current Corp", table_text)

    def test_docx_matches_clean_layout_and_resume_content(self) -> None:
        career_data = self._career().model_dump()
        career_data["skills"] = [
            {"id": "python", "name": "Python", "category": "Languages", "core": True},
            {"id": "yaml", "name": "YAML", "category": "Languages", "level": "advanced"},
        ]
        career_data["experience"][0]["title"] = f"Build {chr(0x2014)} Systems"
        career = CareerDocument.model_validate(career_data)
        with tempfile.TemporaryDirectory() as temporary_directory:
            output_path = write_resume(Path(temporary_directory), career, "en", "docx")
            document = Document(output_path)

        section = document.sections[0]
        texts = [paragraph.text for paragraph in document.paragraphs]
        self.assertEqual(output_path.name, "resume.en.docx")
        self.assertAlmostEqual(section.page_width / 36000, 210, places=1)
        self.assertAlmostEqual(section.page_height / 36000, 297, places=1)
        self.assertAlmostEqual(section.top_margin / 36000, 16, places=1)
        self.assertEqual(document.styles["Normal"].font.name, "Arial")
        self.assertIn("Seriema Example", texts)
        self.assertIn("Professional Experience", texts)
        self.assertIn(f"Build {chr(0x2014)} Systems - Earlier Corp", texts)
        self.assertIn("Jan 2024", "\n".join(texts))
        self.assertIn("Present", "\n".join(texts))
        self.assertIn("Portfolio: https://example.invalid/seriema", texts)
        self.assertIn("Advanced", "\n".join(texts))
        self.assertTrue(any(run.bold for paragraph in document.paragraphs for run in paragraph.runs if run.text == "Python"))
        self.assertEqual(document.tables, [])
        self.assertEqual(len(document.inline_shapes), 0)
        self.assertNotIn("List Bullet", [paragraph.style.name for paragraph in document.paragraphs])
        self.assertIn("- Highlight: Built reliable tools.", texts)
        self.assertNotIn(chr(0xFFFD), "\n".join(texts))
        self.assertTrue(any("w:pBdr" in paragraph._p.xml for paragraph in document.paragraphs if paragraph.text == "Summary"))
        self.assertEqual(document.paragraphs[0].runs[0].font.size, Pt(22))

    def test_docx_localizes_headings_and_omits_empty_sections(self) -> None:
        localized = Document(BytesIO(render_docx(self._career(), "pt-BR")))
        localized_texts = [paragraph.text for paragraph in localized.paragraphs]
        self.assertIn("Resumo", localized_texts)
        self.assertIn("Experiência profissional", localized_texts)
        self.assertIn("Atual", "\n".join(localized_texts))
        self.assertIn("Portuguese | English", localized_texts)

        career = CareerDocument.model_validate(
            {
                "schema_version": 1,
                "profile": {"name": "Seriema", "title": "Engineer", "email": "seriema@example.invalid"},
            }
        )

        document = Document(BytesIO(render_docx(career, "pt-BR")))
        texts = [paragraph.text for paragraph in document.paragraphs]

        self.assertIn("Seriema", texts)
        self.assertNotIn("Resumo", texts)
        self.assertNotIn("ExperiÃªncia profissional", texts)
        self.assertNotIn("Idiomas", texts)

    def test_html_is_semantic_and_escapes_canonical_content(self) -> None:
        career = self._career().model_copy(
            update={"summary": "<script>alert(1)</script>"}
        )

        rendered = render_html(career, "en")

        self.assertIn("<main>", rendered)
        self.assertIn("<section><h2>Summary</h2>", rendered)
        self.assertIn("&lt;script&gt;alert(1)&lt;/script&gt;", rendered)
        self.assertNotIn("<script>alert", rendered)

    def test_pdf_uses_fake_renderer_and_writes_only_pdf(self) -> None:
        class FakePdf:
            def render(self, html: str) -> bytes:
                self.html = html
                return b"%PDF-fake"

        with tempfile.TemporaryDirectory() as temporary_directory:
            fake = FakePdf()
            project_path = Path(temporary_directory)
            output_path = write_resume(project_path, self._career(), "en", "pdf", fake)

            self.assertEqual(output_path.name, "resume.en.pdf")
            self.assertEqual(output_path.read_bytes(), b"%PDF-fake")
            self.assertIn("Professional Experience", fake.html)
            self.assertFalse((project_path / "exports/resume.html").exists())

    def test_style_failure_does_not_overwrite_an_existing_artifact(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            project_path = Path(temporary_directory)
            exports_path = project_path / "exports"
            exports_path.mkdir()
            output_path = exports_path / "resume.md"
            output_path.write_text("previous artifact", encoding="utf-8")

            with self.assertRaisesRegex(ValueError, "Unknown resume style"):
                write_resume(
                    project_path,
                    self._career(),
                    "en",
                    "markdown",
                    style_id="unknown",
                )

            self.assertEqual(output_path.read_text(encoding="utf-8"), "previous artifact")

    def test_renders_localized_headings_and_preserves_user_content(self) -> None:
        career = self._career()

        portuguese = render_markdown(career, "pt-BR")
        english = render_markdown(career, "en")

        self.assertIn("## Resumo", portuguese)
        self.assertIn("## Experiência profissional", portuguese)
        self.assertIn("## Formação acadêmica", portuguese)
        self.assertIn("## Summary", english)
        self.assertIn("## Professional Experience", english)
        self.assertIn("Texto canônico sem tradução.", english)
        self.assertIn("Software Engineer", english)

    def test_omits_empty_sections_and_hides_non_resume_sections(self) -> None:
        career = CareerDocument.model_validate(
            {
                "schema_version": 1,
                "profile": {
                    "name": "Seriema Example",
                    "title": "Engineer",
                    "email": "seriema@example.invalid",
                },
                "evidence": [{"id": "private-proof", "statement": "Do not render"}],
            }
        )

        rendered = render_markdown(career, "en")

        self.assertNotIn("## Summary", rendered)
        self.assertNotIn("## Professional Experience", rendered)
        self.assertNotIn("## Education", rendered)
        self.assertNotIn("## Skills", rendered)
        self.assertNotIn("## Languages", rendered)
        self.assertNotIn("Do not render", rendered)

    def test_orders_dated_sections_in_reverse_chronology_and_keeps_skill_order(
        self,
    ) -> None:
        rendered = render_markdown(self._career(), "en")

        self.assertLess(rendered.index("Current Corp"), rendered.index("Earlier Corp"))
        self.assertLess(
            rendered.index("New University"), rendered.index("Old University")
        )
        self.assertLess(rendered.index("Python"), rendered.index("YAML"))
        self.assertIn("Jan 2024 - Present", rendered)
        self.assertIn("Remote | Contract", rendered)

    def test_writes_the_single_resume_artifact(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            project_path = Path(temporary_directory)
            exports_path = project_path / "exports"
            exports_path.mkdir()
            output_path = write_markdown_resume(project_path, self._career(), "pt-BR")

            self.assertEqual(output_path, exports_path / "resume.pt-BR.md")
            self.assertTrue(output_path.exists())

    def test_atomic_write_keeps_existing_artifact_when_replacement_fails(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            project_path = Path(temporary_directory)
            exports_path = project_path / "exports"
            exports_path.mkdir()
            output_path = exports_path / "resume.md"
            output_path.write_text("previous artifact", encoding="utf-8")

            with patch(
                "seriemacv.renderer.os.replace", side_effect=OSError("disk error")
            ):
                with self.assertRaises(OSError):
                    write_markdown_resume(project_path, self._career(), "en")

            self.assertEqual(
                output_path.read_text(encoding="utf-8"), "previous artifact"
            )
            self.assertEqual(list(exports_path.glob(".resume.md.*.tmp")), [])

    def test_docx_atomic_write_keeps_existing_artifact_when_replacement_fails(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            project_path = Path(temporary_directory)
            exports_path = project_path / "exports"
            exports_path.mkdir()
            output_path = exports_path / "resume.docx"
            output_path.write_bytes(b"previous artifact")

            with patch(
                "seriemacv.renderer.os.replace", side_effect=OSError("disk error")
            ):
                with self.assertRaises(OSError):
                    write_resume(project_path, self._career(), "en", "docx")

            self.assertEqual(output_path.read_bytes(), b"previous artifact")
            self.assertEqual(list(exports_path.glob(".resume.docx.*.tmp")), [])

    @staticmethod
    def _career() -> CareerDocument:
        return CareerDocument.model_validate(
            {
                "schema_version": 1,
                "profile": {
                    "name": "Seriema Example",
                    "title": "Software Engineer",
                    "location": "Remote",
                    "email": "seriema@example.invalid",
                    "phone": "+1 555 0100",
                    "links": {"Portfolio": "https://example.invalid/seriema"},
                    "languages": ["Portuguese", "English"],
                },
                "summary": "Texto canônico sem tradução.",
                "experience": [
                    {
                        "id": "earlier-corp",
                        "company": "Earlier Corp",
                        "title": "Developer",
                        "start_date": "2020-01",
                        "end_date": "2023-12",
                    },
                    {
                        "id": "current-corp",
                        "company": "Current Corp",
                        "title": "Software Engineer",
                        "start_date": "2024-01",
                        "location": "Remote",
                        "employment_type": "Contract",
                        "highlights": ["Built reliable tools."],
                    },
                ],
                "education": [
                    {
                        "id": "old-university",
                        "institution": "Old University",
                        "degree": "BSc",
                        "start_date": "2014-01",
                        "end_date": "2018-12",
                    },
                    {
                        "id": "new-university",
                        "institution": "New University",
                        "degree": "MSc",
                        "start_date": "2021-01",
                        "end_date": "2023-12",
                    },
                ],
                "skills": [
                    {"id": "python", "name": "Python"},
                    {"id": "yaml", "name": "YAML"},
                ],
            }
        )


class ResumeRenderCliTests(unittest.TestCase):
    def test_cli_uses_configured_style_and_allows_a_one_off_override(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            project_path = Path(temporary_directory) / "career-project"
            create_project(
                project_path,
                project_name="Career Project",
                resume_language="en",
                resume_style="modern",
            )
            self._write_complete_career(project_path)

            with redirect_stdout(StringIO()):
                result = main([
                    "resume", "render", str(project_path), "--format", "html",
                ])
            self.assertEqual(result, 0)
            output_path = project_path / "exports/resume.en.html"
            self.assertIn('data-style="modern"', output_path.read_text(encoding="utf-8"))
            self.assertIn("#647d74", output_path.read_text(encoding="utf-8"))

            with redirect_stdout(StringIO()):
                result = main([
                    "resume", "render", str(project_path), "--format", "html",
                    "--style", "compact-alt",
                ])
            self.assertEqual(result, 0)
            self.assertIn('data-style="compact-alt"', output_path.read_text(encoding="utf-8"))
            self.assertIn(
                "resume_style: modern",
                (project_path / "seriemacv.yml").read_text(encoding="utf-8"),
            )

    def test_cli_uses_project_resume_language_and_writes_single_artifact(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            project_path = Path(temporary_directory) / "career-project"
            create_project(
                project_path, project_name="Career Project", resume_language="en"
            )
            self._write_complete_career(project_path)

            result = main(
                [
                    "resume",
                    "render",
                    str(project_path),
                    "--format",
                    "markdown",
                ]
            )

            self.assertEqual(result, 0)
            output_path = project_path / "exports/resume.en.md"
            self.assertTrue(output_path.exists())
            self.assertIn("## Summary", output_path.read_text(encoding="utf-8"))

            result = main([
                "resume", "render", str(project_path), "--format", "html",
            ])

            self.assertEqual(result, 0)
            self.assertTrue((project_path / "exports/resume.en.html").exists())

            result = main([
                "resume", "render", str(project_path), "--format", "docx",
            ])

            self.assertEqual(result, 0)
            self.assertTrue((project_path / "exports/resume.en.docx").exists())

    def test_cli_does_not_overwrite_artifact_when_career_is_incomplete(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            project_path = Path(temporary_directory) / "career-project"
            create_project(project_path, project_name="Career Project")
            output_path = project_path / "exports/resume.md"
            output_path.write_text("previous artifact", encoding="utf-8")
            stderr = StringIO()

            with redirect_stderr(stderr):
                result = main(
                    [
                        "resume",
                        "render",
                        str(project_path),
                        "--format",
                        "markdown",
                    ]
                )

            self.assertEqual(result, 1)
            self.assertIn("profile.name", stderr.getvalue())
            self.assertEqual(
                output_path.read_text(encoding="utf-8"), "previous artifact"
            )

    def test_cli_does_not_overwrite_docx_when_career_is_incomplete(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            project_path = Path(temporary_directory) / "career-project"
            create_project(project_path, project_name="Career Project")
            output_path = project_path / "exports/resume.docx"
            output_path.write_bytes(b"previous DOCX artifact")

            with redirect_stderr(StringIO()):
                result = main([
                    "resume", "render", str(project_path), "--format", "docx",
                ])

            self.assertEqual(result, 1)
            self.assertEqual(output_path.read_bytes(), b"previous DOCX artifact")

    def test_cli_reports_invalid_project_configuration_at_its_path(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            project_path = Path(temporary_directory) / "career-project"
            create_project(project_path, project_name="Career Project")
            (project_path / "career.yml").write_text(
                """schema_version: 2
profile:
  name: Seriema Example
  email: seriema@example.invalid
experience: []
education: []
skills: []
evidence: []
answers: []
stories: []
""",
                encoding="utf-8",
            )
            configuration_path = project_path / "seriemacv.yml"
            configuration_path.write_text(
                "schema_version: 2\nproject_name: Career Project\nresume_language: invalid locale\n",
                encoding="utf-8",
            )
            stderr = StringIO()

            with redirect_stderr(stderr):
                result = main(
                    [
                        "resume",
                        "render",
                        str(project_path),
                        "--format",
                        "markdown",
                    ]
                )

            self.assertEqual(result, 1)
            self.assertIn(str(configuration_path), stderr.getvalue())
            self.assertFalse((project_path / "exports/resume.pt-BR.md").exists())

    def test_cli_does_not_overwrite_artifact_for_invalid_resume_color(self) -> None:
        with tempfile.TemporaryDirectory() as temporary_directory:
            project_path = Path(temporary_directory) / "career-project"
            create_project(project_path, project_name="Career Project", resume_language="en")
            self._write_complete_career(project_path)
            output_path = project_path / "exports/resume.en.html"
            output_path.write_text("previous artifact", encoding="utf-8")
            (project_path / "seriemacv.yml").write_text(
                "schema_version: 2\nproject_name: Career Project\nresume_language: en\nresume_color: invalid\n",
                encoding="utf-8",
            )

            with redirect_stderr(StringIO()):
                result = main(["resume", "render", str(project_path), "--format", "html"])

            self.assertEqual(result, 1)
            self.assertEqual(output_path.read_text(encoding="utf-8"), "previous artifact")

    @staticmethod
    def _write_complete_career(project_path: Path) -> None:
        (project_path / "career.yml").write_text(
            """schema_version: 2
profile:
  name: Seriema Example
  email: seriema@example.invalid
experience: []
education: []
skills: []
evidence: []
answers: []
stories: []
""",
            encoding="utf-8",
        )
        (project_path / "career.locales" / "en.yml").write_text(
            """schema_version: 1
locale: en
profile: {title: Engineer}
summary: Canonical content.
experience: {}
education: {}
skills: {}
""",
            encoding="utf-8",
        )
